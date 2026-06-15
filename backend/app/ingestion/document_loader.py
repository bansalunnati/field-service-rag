"""
ingestion/document_loader.py

Loads PDF, DOCX, TXT, MD, and CSV files from a directory or as single files.
Auto-detects the correct pipeline from the filename.

Pipeline keys (updated in Phase 1):
  equipment    → technical manuals, specs, maintenance guides
  safety       → SOPs, policies, compliance docs, regulations
  field_reports → submitted reports, general documents
"""

import os
from typing import List, Optional
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader


def load_documents(
    data_path: str = "data/policies",
    pipeline: Optional[str] = None,
) -> List[Document]:
    """
    Loads all supported files from a directory.
    Stamps each document with source + pipeline metadata for citations.

    Args:
        data_path : Directory containing documents to ingest.
        pipeline  : Override auto-detection and force all files into this pipeline.

    Returns:
        Flat list of LangChain Document objects.
    """
    documents = []
    files = os.listdir(data_path)

    print(f"\nFound {len(files)} files in '{data_path}'\n")

    for file in files:
        file_path = os.path.join(data_path, file)
        ext = os.path.splitext(file)[1].lower()
        detected_pipeline = pipeline or _detect_pipeline(file)

        print(f"Loading: {file}  →  pipeline: {detected_pipeline}")

        try:
            if ext == ".pdf":
                docs = _load_pdf(file_path)
            elif ext == ".docx":
                docs = _load_docx(file_path)
            elif ext in (".txt", ".md"):
                docs = _load_txt(file_path)
            elif ext == ".csv":
                docs = _load_csv(file_path)
            else:
                print(f"  Skipping unsupported type: {ext}")
                continue

            for doc in docs:
                doc.metadata.setdefault("source", file)
                doc.metadata["pipeline"] = detected_pipeline

            print(f"  Loaded {len(docs)} pages")
            documents.extend(docs)

        except Exception as e:
            print(f"  ERROR loading {file}: {e}")

    return documents


# ── Individual loaders ────────────────────────────────────────────────────────

def _load_pdf(path: str) -> List[Document]:
    loader = PyPDFLoader(path)
    return loader.load()


def _load_docx(path: str) -> List[Document]:
    try:
        from langchain_community.document_loaders import Docx2txtLoader
        return Docx2txtLoader(path).load()
    except ImportError:
        from docx import Document as DocxDoc
        doc = DocxDoc(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(page_content=text, metadata={"source": os.path.basename(path)})]


def _load_txt(path: str) -> List[Document]:
    from langchain_community.document_loaders import TextLoader
    return TextLoader(path, encoding="utf-8").load()


def _load_csv(path: str) -> List[Document]:
    import csv
    rows = []
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        reader = csv.DictReader(f)
        for row in reader:
            line = ", ".join(f"{k}: {v}" for k, v in row.items() if v)
            rows.append(line)

    docs = []
    for i in range(0, len(rows), 50):
        block = "\n".join(rows[i: i + 50])
        docs.append(Document(
            page_content=block,
            metadata={"source": os.path.basename(path), "page": i // 50 + 1},
        ))
    return docs


# ── Pipeline auto-detection ───────────────────────────────────────────────────

def _detect_pipeline(filename: str) -> str:
    name = filename.lower()

    # Hardcoded overrides for known seed files
    hardcoded = {
        "equipment and rules compliance.pdf": "equipment",
        "hazardous materials field inspection checklist.pdf": "safety",
        "sop telecom towers.pdf": "safety",
    }
    if name in hardcoded:
        return hardcoded[name]

    # Keyword fallback for any future files
    safety_keywords = [
        "hazardous", "inspection", "checklist", "sop", "procedure",
        "regulation", "emergency", "protocol", "audit",
    ]
    equipment_keywords = [
        "manual", "spec", "part", "maintenance", "asset",
        "machine", "tool", "hardware", "calibration", "installation",
    ]

    if any(k in name for k in safety_keywords):
        return "safety"
    if any(k in name for k in equipment_keywords):
        return "equipment"

    return "field_reports"